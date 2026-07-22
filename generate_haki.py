from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ===== COVER PAGE =====
for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('APLIKASI BOOKING BARBERSHOP')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('DENGAN SISTEM PEMBAYARAN QRIS DAN VIRTUAL ACCOUNT')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('BERBASIS MOBILE DAN WEB')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')

# Info table
info_data = [
    ('Platform', 'Mobile (Android/iOS) & Web Admin'),
    ('Teknologi', 'Laravel (Backend API) + Flutter (Frontend Mobile)'),
    ('Payment Gateway', 'Pakasir — QRIS, Virtual Account (BNI, BRI, CIMB, Maybank, Permata, BNC, ATM Bersama, Sampoerna, Artha Graha), PayPal'),
    ('Database', 'PostgreSQL (Neon)'),
    ('Hosting', 'Render (Docker) + Neon PostgreSQL'),
]

table = doc.add_table(rows=len(info_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(info_data):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.style = doc.styles['Normal']
            for run in paragraph.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
    row.cells[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph('')
doc.add_paragraph('')

# Pemegang Hak Cipta
hak_cipta = doc.add_paragraph()
hak_cipta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hak_cipta.add_run('PEMEGANG HAK CIPTA')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph('')

hak_table = doc.add_table(rows=2, cols=5)
hak_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['No', 'Nama', 'Email', 'No. Telepon', 'Kode Pos']
for i, h in enumerate(headers):
    hak_table.rows[0].cells[i].text = h
    for p in hak_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'

hak_table.rows[1].cells[0].text = '1'
hak_table.rows[1].cells[1].text = 'Irya Muhammad Ghiffari'
hak_table.rows[1].cells[2].text = '[email Lo]'
hak_table.rows[1].cells[3].text = '[no HP Lo]'
hak_table.rows[1].cells[4].text = '[kode pos Lo]'
for cell in hak_table.rows[1].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'

doc.add_paragraph('')
doc.add_paragraph('')

tahun = doc.add_paragraph()
tahun.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tahun.add_run('TAHUN PENDAFTARAN: 2025')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

# Page break
doc.add_page_break()

# ===== COVER - SCREENSHOTS =====
cover_heading = doc.add_heading('COVER — TAMPILAN APLIKASI', level=1)
for run in cover_heading.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph('')

# Cover section
cover_sub = doc.add_heading('Cover (Tampilan Utama)', level=2)
for run in cover_sub.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph('[Tempel Gambar 1 — Home Screen di sini]')
doc.add_paragraph('[Tempel Gambar 2 — Login/Register di sini]')

doc.add_paragraph('')

# Pelanggan section
pelanggan_heading = doc.add_heading('Pelanggan (Flow Booking & Pembayaran)', level=2)
for run in pelanggan_heading.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph('[Tempel Gambar 3 — Pilih Layanan & Barber]')
doc.add_paragraph('[Tempel Gambar 4 — Pilih Jadwal]')
doc.add_paragraph('[Tempel Gambar 5 — Konfirmasi Booking]')
doc.add_paragraph('[Tempel Gambar 6 — Pembayaran (QRIS/VA/Tunai)]')
doc.add_paragraph('[Tempel Gambar 7 — Riwayat Booking]')

doc.add_paragraph('')

# Admin section
admin_heading = doc.add_heading('Admin (Dashboard & Manajemen)', level=2)
for run in admin_heading.runs:
    run.font.name = 'Times New Roman'

doc.add_paragraph('[Tempel Gambar 8 — Dashboard Admin]')
doc.add_paragraph('[Tempel Gambar 9 — Kelola Booking]')
doc.add_paragraph('[Tempel Gambar 10 — Transaksi & Laporan]')

doc.add_page_break()

# ===== SCREENSHOT LIST =====
ss_heading = doc.add_heading('DAFTAR SCREENSHOT YANG HARUS DISIAPKAN', level=1)
for run in ss_heading.runs:
    run.font.name = 'Times New Roman'

ss_data = [
    ('01_home.png', 'Home — daftar layanan & barber'),
    ('02_login.png', 'Login / Register (OTP)'),
    ('03_layanan.png', 'Pilih layanan & barber'),
    ('04_jadwal.png', 'Pilih tanggal & jam'),
    ('05_konfirmasi.png', 'Booking dikonfirmasi admin'),
    ('06_bayar.png', 'Halaman bayar (QRIS/VA/Tunai)'),
    ('07_riwayat.png', 'Riwayat booking pelanggan'),
    ('08_admin_dash.png', 'Dashboard admin (statistik)'),
    ('09_kelola_booking.png', 'Kelola booking (konfirmasi/batal)'),
    ('10_transaksi.png', 'Transaksi & laporan pendapatan'),
]

ss_table = doc.add_table(rows=len(ss_data)+1, cols=3)
ss_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ss_table.rows[0].cells[0].text = 'No'
ss_table.rows[0].cells[1].text = 'File'
ss_table.rows[0].cells[2].text = 'Deskripsi'
for cell in ss_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'

for i, (file, desc) in enumerate(ss_data):
    row = ss_table.rows[i+1]
    row.cells[0].text = str(i+1)
    row.cells[1].text = file
    row.cells[2].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'

doc.add_page_break()

# ===== DESKRIPSI =====
desc_heading = doc.add_heading('DESKRIPSI APLIKASI', level=1)
for run in desc_heading.runs:
    run.font.name = 'Times New Roman'

desc_text = (
    'Aplikasi Booking Barbershop yang digunakan oleh Pelanggan untuk memilih layanan, '
    'barber, dan jadwal secara online. Dilengkapi sistem pembayaran melalui QRIS dan '
    'Virtual Account serta pembayaran tunai di tempat. Memberikan kemudahan booking, '
    'riwayat transaksi, dan manajemen booking bagi admin. Dibangun menggunakan Laravel '
    'sebagai backend API dan Flutter sebagai frontend mobile.'
)
p = doc.add_paragraph(desc_text)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ===== SOURCE CODE =====
code_heading = doc.add_heading('SOURCE CODE', level=1)
for run in code_heading.runs:
    run.font.name = 'Times New Roman'

# Potongan 1
code1_heading = doc.add_heading('Potongan 1 — Backend Integrasi Pakasir', level=2)
for run in code1_heading.runs:
    run.font.name = 'Times New Roman'

code1 = '''// Membuat transaksi pembayaran melalui Payment Gateway Pakasir
$booking = Booking::with(['user', 'services'])->findOrFail($bookingId);
$orderId = 'ARF-PAY-' . $booking->id . '-' . time();

$response = Http::post('https://app.pakasir.com/api/transactioncreate/all', [
    'project' => 'barbershop',
    'order_id' => $orderId,
    'amount' => (int) $booking->total_price,
    'api_key' => config('pakasir.api_key'),
]);

$paymentData = $response->json()['payment'];
$payment = Payment::create([
    'booking_id' => $booking->id,
    'order_id' => $orderId,
    'amount' => $booking->total_price,
    'status' => 'pending',
    'payment_url' => $paymentData['payment_url'],
    'payment_method' => $paymentData['payment_method'] ?? 'qris',
]);

return response()->json([
    'success' => true,
    'data' => ['payment_url' => $payment->payment_url]
]);'''

p = doc.add_paragraph(code1)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_paragraph('')

# Potongan 2
code2_heading = doc.add_heading('Potongan 2 — WebView Pembayaran Flutter', level=2)
for run in code2_heading.runs:
    run.font.name = 'Times New Roman'

code2 = '''// WebView pembayaran dengan polling status
class PakasirPaymentWebView extends StatefulWidget {
  final int bookingId;
  final String initialPaymentUrl;
  final VoidCallback onClose;
  final void Function(bool success) onPaymentFinished;

  State<PakasirPaymentWebView> createState() => _State();
}

class _State extends State<PakasirPaymentWebView> {
  late WebViewController _ctrl;
  Timer? _pollTimer;

  void initState() {
    super.initState();
    _ctrl = WebViewController()
      ..setJavaScriptMode(JavaScriptMode.unrestricted)
      ..loadRequest(Uri.parse(widget.initialPaymentUrl));
    _pollTimer = Timer.periodic(Duration(seconds: 3), (_) => _verifyPaid());
  }

  Future<void> _verifyPaid() async {
    if (!mounted) return;
    final paid = await context
        .read<BookingProvider>()
        .verifyBookingPaid(widget.bookingId);
    if (paid) {
      _pollTimer?.cancel();
      widget.onPaymentFinished(true);
    }
  }

  Widget build(BuildContext context) {
    return Scaffold(
      appBar: AppBar(title: Text('Pembayaran')),
      body: WebViewWidget(controller: _ctrl),
    );
  }
}'''

p = doc.add_paragraph(code2)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_paragraph('')

# Potongan 3
code3_heading = doc.add_heading('Potongan 3 — Model + Logika Bisnis Booking', level=2)
for run in code3_heading.runs:
    run.font.name = 'Times New Roman'

code3 = '''// Model dan logika booking
class BookingModel {
  final int id;
  final String bookingCode;
  final String bookingDate;
  final String bookingTime;
  final double totalPrice;
  final String status;
  final PaymentModel? payment;

  BookingModel({
    required this.id,
    required this.bookingCode,
    required this.bookingDate,
    required this.bookingTime,
    required this.totalPrice,
    required this.status,
    this.payment,
  });

  factory BookingModel.fromJson(Map<String, dynamic> json) {
    return BookingModel(
      id: json['id'],
      bookingCode: json['booking_code'],
      bookingDate: json['booking_date'],
      bookingTime: json['booking_time'],
      totalPrice: double.parse(json['total_price'].toString()),
      status: json['status'],
      payment: json['payment'] != null
          ? PaymentModel.fromJson(json['payment'])
          : null,
    );
  }

  bool get isPaid => payment?.status == 'paid';
  bool get canPayGateway => status == 'confirmed' && !isPaid;
  bool get canReschedule => (status == 'pending' || status == 'confirmed') && !isPaid;
  bool get canCancel => (status == 'pending' || status == 'confirmed') && !isPaid;

  String get priceFormatted {
    final formatted = totalPrice.toStringAsFixed(0).replaceAllMapped(
        RegExp(r'(\d{1,3})(?=(\d{3})+(?!\d))'),
        (m) => '${m[1]}.');
    return 'Rp $formatted';
  }
}'''

p = doc.add_paragraph(code3)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_page_break()

# ===== MANUAL BOOK =====
manual_heading = doc.add_heading('MANUAL BOOK / BUKU PANDUAN', level=1)
for run in manual_heading.runs:
    run.font.name = 'Times New Roman'

# Pelanggan
pelanggan_manual = doc.add_heading('Untuk Pelanggan', level=2)
for run in pelanggan_manual.runs:
    run.font.name = 'Times New Roman'

pelanggan_steps = [
    'Login/Register — Masuk atau daftar akun baru',
    'Pilih Layanan & Barber — Pilih layanan potong rambut, barber favorit',
    'Pilih Jadwal — Tentukan tanggal & jam, kirim booking',
    'Tunggu Konfirmasi Admin — Notifikasi muncul saat dikonfirmasi',
    'Pembayaran — Pilih "Bayar Online (QRIS/VA)" atau "Bayar di Tempat"',
    'Riwayat Booking — Lihat status booking (Menunggu, Dikonfirmasi, Selesai, Dibatalkan)',
]

for i, step in enumerate(pelanggan_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

doc.add_paragraph('')

# Admin
admin_manual = doc.add_heading('Untuk Admin', level=2)
for run in admin_manual.runs:
    run.font.name = 'Times New Roman'

admin_steps = [
    'Login Admin — Masuk dengan akun admin',
    'Kelola Booking — Konfirmasi, batalkan, ubah status booking',
    'Kelola Layanan — Tambah/edit/hapus layanan & harga',
    'Kelola Barber — Tambah/edit/hapus barber, jadwal tidak tersedia',
    'Transaksi — Lihat riwayat pembayaran, filter by status/tanggal',
    'Laporan Pendapatan — Lihat rekap pendapatan per tanggal',
]

for i, step in enumerate(admin_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.style = doc.styles['Normal']
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

# Save
output_path = r'D:\FILE IRYA\SEMESTER 8\SKRIPSI\Semhas\arfan_barbershop\HAKI_Arfan_Barbershop.docx'
doc.save(output_path)
print(f'File saved to: {output_path}')