from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ===== PAGE SETUP =====
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ===== STYLES =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0, 51, 102)
    if level == 1:
        h_style.font.size = Pt(16)
        h_style.font.bold = True
    elif level == 2:
        h_style.font.size = Pt(14)
        h_style.font.bold = True
    else:
        h_style.font.size = Pt(12)
        h_style.font.bold = True

# Code style
if 'Code' not in doc.styles:
    code_style = doc.styles.add_style('Code', 1)
else:
    code_style = doc.styles['Code']
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9.5)
code_style.font.color.rgb = RGBColor(30, 30, 30)
pf = code_style.paragraph_format
pf.space_before = Pt(4)
pf.space_after = Pt(4)
pf.line_spacing = 1.0
pf.left_indent = Cm(1)

# Helper functions
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text):
    """Add formatted code block with light gray background"""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(style='Code')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
    return p

def format_table(table, header_color='003366'):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    # Data rows
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    # Alternating row colors
    for i, row in enumerate(table.rows[1:], 1):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, 'F2F7FB')

def add_placeholder_image(doc, text, width=Cm(7)):
    """Add a styled placeholder for screenshot"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True
    run.font.name = 'Times New Roman'
    # Add border-like spacing
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# ============================================
# DOCUMENT START
# ============================================
doc = Document()

# Re-apply page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# ===== COVER PAGE =====
for _ in range(5):
    doc.add_paragraph('')

# Title block
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run('APLIKASI BOOKING BARBERSHOP')
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
run = subtitle.add_run('DENGAN SISTEM PEMBAYARAN QRIS DAN VIRTUAL ACCOUNT')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.paragraph_format.space_after = Pt(24)
run = subtitle2.add_run('BERBASIS MOBILE DAN WEB')
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

# Horizontal line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 60)
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.size = Pt(12)

doc.add_paragraph('')

# Info table
info_data = [
    ('Platform', 'Mobile (Android/iOS) & Web Admin'),
    ('Teknologi', 'Laravel (Backend API) + Flutter (Frontend Mobile)'),
    ('Payment Gateway', 'Pakasir — QRIS, Virtual Account (BNI, BRI, CIMB, Maybank, Permata, BNC, ATM Bersama, Sampoerna, Artha Graha), PayPal'),
    ('Database', 'PostgreSQL (Neon)'),
    ('Hosting', 'Render (Docker) + Neon PostgreSQL'),
]

table = doc.add_table(rows=len(info_data), cols=2)
table.autofit = True
for i, (label, value) in enumerate(info_data):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(12)

format_table(table)

doc.add_paragraph('')
doc.add_paragraph('')

# Pemegang Hak Cipta
hak_title = doc.add_paragraph()
hak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
hak_title.paragraph_format.space_before = Pt(12)
hak_title.paragraph_format.space_after = Pt(8)
run = hak_title.add_run('PEMEGANG HAK CIPTA')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

hak_table = doc.add_table(rows=2, cols=5)
hak_table.autofit = True
headers = ['No', 'Nama', 'Email', 'No. Telepon', 'Kode Pos']
for i, h in enumerate(headers):
    hak_table.rows[0].cells[i].text = h

hak_table.rows[1].cells[0].text = '1'
hak_table.rows[1].cells[1].text = 'Irya Muhammad Ghiffari'
hak_table.rows[1].cells[2].text = '[email Lo]'
hak_table.rows[1].cells[3].text = '[no HP Lo]'
hak_table.rows[1].cells[4].text = '[kode pos Lo]'

format_table(hak_table)

doc.add_paragraph('')
doc.add_paragraph('')

tahun = doc.add_paragraph()
tahun.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tahun.add_run('TAHUN PENDAFTARAN: 2025')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0, 51, 102)

# Page break
doc.add_page_break()

# ===== COVER - SCREENSHOTS =====
heading = doc.add_heading('COVER — TAMPILAN APLIKASI', level=1)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')

# Cover section
sub = doc.add_heading('Cover (Tampilan Utama)', level=2)
add_placeholder_image(doc, '[ Tempel Gambar 1 — Home Screen ]')
add_placeholder_image(doc, '[ Tempel Gambar 2 — Login / Register ]')

doc.add_paragraph('')

# Pelanggan section
sub = doc.add_heading('Pelanggan (Flow Booking & Pembayaran)', level=2)
screenshots_pelanggan = [
    'Pilih Layanan & Barber',
    'Pilih Jadwal',
    'Konfirmasi Booking',
    'Pembayaran (QRIS / VA / Tunai)',
    'Riwayat Booking',
]
for s in screenshots_pelanggan:
    add_placeholder_image(doc, f'[ Tempel Gambar — {s} ]')

doc.add_paragraph('')

# Admin section
sub = doc.add_heading('Admin (Dashboard & Manajemen)', level=2)
screenshots_admin = [
    'Dashboard Admin',
    'Kelola Booking',
    'Transaksi & Laporan Pendapatan',
]
for s in screenshots_admin:
    add_placeholder_image(doc, f'[ Tempel Gambar — {s} ]')

doc.add_page_break()

# ===== SCREENSHOT LIST =====
heading = doc.add_heading('DAFTAR SCREENSHOT YANG HARUS DISIAPKAN', level=1)

ss_data = [
    ('01_home.png', 'Home — daftar layanan & barber'),
    ('02_login.png', 'Login / Register (OTP)'),
    ('03_layanan.png', 'Pilih layanan & barber'),
    ('04_jadwal.png', 'Pilih tanggal & jam'),
    ('05_konfirmasi.png', 'Booking dikonfirmasi admin'),
    ('06_bayar.png', 'Halaman bayar (QRIS/VA/Tunai)'),
    ('07_riwayat.png', 'Riwayat booking pelanggan'),
    ('08_admin_dash.png', 'Dashboard admin (statistik)'),
    ('09_kelola_booking.png', 'Kelola booking (konfirmasi/batal)'),
    ('10_transaksi.png', 'Transaksi & laporan pendapatan'),
]

ss_table = doc.add_table(rows=len(ss_data)+1, cols=3)
ss_table.autofit = True
ss_table.rows[0].cells[0].text = 'No'
ss_table.rows[0].cells[1].text = 'Nama File'
ss_table.rows[0].cells[2].text = 'Deskripsi'

for i, (file, desc) in enumerate(ss_data):
    row = ss_table.rows[i+1]
    row.cells[0].text = str(i+1)
    row.cells[1].text = file
    row.cells[2].text = desc

format_table(ss_table)

doc.add_page_break()

# ===== DESKRIPSI =====
heading = doc.add_heading('DESKRIPSI APLIKASI', level=1)

desc_text = (
    'Aplikasi Booking Barbershop yang digunakan oleh Pelanggan untuk memilih layanan, '
    'barber, dan jadwal secara online. Dilengkapi sistem pembayaran melalui QRIS dan '
    'Virtual Account serta pembayaran tunai di tempat. Memberikan kemudahan booking, '
    'riwayat transaksi, dan manajemen booking bagi admin. Dibangun menggunakan Laravel '
    'sebagai backend API dan Flutter sebagai frontend mobile.'
)
p = doc.add_paragraph(desc_text)
p.paragraph_format.first_line_indent = Cm(1)

doc.add_page_break()

# ===== SOURCE CODE =====
heading = doc.add_heading('SOURCE CODE', level=1)

# Potongan 1
sub = doc.add_heading('Potongan 1 — Backend Integrasi Pakasir', level=2)
p = doc.add_paragraph('Kode ini menangani pembuatan transaksi pembayaran melalui Payment Gateway Pakasir. '
    'Backend Laravel mengirim request ke API Pakasir, menerima URL pembayaran, '
    'dan menyimpannya ke database.')
p.paragraph_format.space_after = Pt(6)

code1 = '''// Membuat transaksi pembayaran melalui Payment Gateway Pakasir
$booking = Booking::with(['user', 'services'])->findOrFail($bookingId);
$orderId = 'ARF-PAY-' . $booking->id . '-' . time();

$response = Http::post('https://app.pakasir.com/api/transactioncreate/all', [
    'project' => 'barbershop',
    'order_id' => $orderId,
    'amount' => (int) $booking->total_price,
    'api_key' => config('pakasir.api_key'),
]);

$paymentData = $response->json()['payment'];
$payment = Payment::create([
    'booking_id' => $booking->id,
    'order_id' => $orderId,
    'amount' => $booking->total_price,
    'status' => 'pending',
    'payment_url' => $paymentData['payment_url'],
    'payment_method' => $paymentData['payment_method'] ?? 'qris',
]);

return response()->json([
    'success' => true,
    'data' => ['payment_url' => $payment->payment_url]
]);'''

add_code_block(doc, code1)

doc.add_paragraph('')

# Potongan 2
sub = doc.add_heading('Potongan 2 — WebView Pembayaran Flutter', level=2)
p = doc.add_paragraph('Kode Flutter ini menampilkan halaman pembayaran Pakasir di dalam WebView. '
    'Setiap 3 detik, aplikasi memeriksa status pembayaran ke backend melalui polling. '
    'Jika pembayaran berhasil, WebView ditutup otomatis.')
p.paragraph_format.space_after = Pt(6)

code2 = '''// WebView pembayaran dengan polling status
class PakasirPaymentWebView extends StatefulWidget {
  final int bookingId;
  final String initialPaymentUrl;
  final VoidCallback onClose;
  final void Function(bool success) onPaymentFinished;

  State<PakasirPaymentWebView> createState() => _State();
}

class _State extends State<PakasirPaymentWebView> {
  late WebViewController _ctrl;
  Timer? _pollTimer;

  void initState() {
    super.initState();
    _ctrl = WebViewController()
      ..setJavaScriptMode(JavaScriptMode.unrestricted)
      ..loadRequest(Uri.parse(widget.initialPaymentUrl));
    _pollTimer = Timer.periodic(Duration(seconds: 3), (_) => _verifyPaid());
  }

  Future<void> _verifyPaid() async {
    if (!mounted) return;
    final paid = await context
        .read<BookingProvider>()
        .verifyBookingPaid(widget.bookingId);
    if (paid) {
      _pollTimer?.cancel();
      widget.onPaymentFinished(true);
    }
  }

  Widget build(BuildContext context) {
    return Scaffold(
      appBar: AppBar(title: Text('Pembayaran')),
      body: WebViewWidget(controller: _ctrl),
    );
  }
}'''

add_code_block(doc, code2)

doc.add_paragraph('')

# Potongan 3
sub = doc.add_heading('Potongan 3 — Model + Logika Bisnis Booking', level=2)
p = doc.add_paragraph('Model data booking yang merepresentasikan informasi pemesanan, '
    'termasuk status pembayaran dan method untuk menentukan aksi yang diizinkan '
    '(bayar, ubah jadwal, batalkan).')
p.paragraph_format.space_after = Pt(6)

code3 = '''// Model dan logika booking
class BookingModel {
  final int id;
  final String bookingCode;
  final String bookingDate;
  final String bookingTime;
  final double totalPrice;
  final String status;
  final PaymentModel? payment;

  BookingModel({
    required this.id,
    required this.bookingCode,
    required this.bookingDate,
    required this.bookingTime,
    required this.totalPrice,
    required this.status,
    this.payment,
  });

  factory BookingModel.fromJson(Map<String, dynamic> json) {
    return BookingModel(
      id: json['id'],
      bookingCode: json['booking_code'],
      bookingDate: json['booking_date'],
      bookingTime: json['booking_time'],
      totalPrice: double.parse(json['total_price'].toString()),
      status: json['status'],
      payment: json['payment'] != null
          ? PaymentModel.fromJson(json['payment'])
          : null,
    );
  }

  bool get isPaid => payment?.status == 'paid';
  bool get canPayGateway => status == 'confirmed' && !isPaid;
  bool get canReschedule => (status == 'pending' || status == 'confirmed') && !isPaid;
  bool get canCancel => (status == 'pending' || status == 'confirmed') && !isPaid;

  String get priceFormatted {
    final formatted = totalPrice.toStringAsFixed(0).replaceAllMapped(
        RegExp(r'(\\d{1,3})(?=(\\d{3})+(?!\\d))'),
        (m) => '${m[1]}.');
    return 'Rp \$formatted';
  }
}'''

add_code_block(doc, code3)

doc.add_page_break()

# ===== MANUAL BOOK =====
heading = doc.add_heading('MANUAL BOOK / BUKU PANDUAN', level=1)

# Pelanggan
sub = doc.add_heading('Untuk Pelanggan', level=2)
pelanggan_steps = [
    'Login/Register — Masuk atau daftar akun baru',
    'Pilih Layanan & Barber — Pilih layanan potong rambut, barber favorit',
    'Pilih Jadwal — Tentukan tanggal & jam, kirim booking',
    'Tunggu Konfirmasi Admin — Notifikasi muncul saat dikonfirmasi',
    'Pembayaran — Pilih "Bayar Online (QRIS/VA)" atau "Bayar di Tempat"',
    'Riwayat Booking — Lihat status booking (Menunggu, Dikonfirmasi, Selesai, Dibatalkan)',
]

for i, step in enumerate(pelanggan_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph('')

# Admin
sub = doc.add_heading('Untuk Admin', level=2)
admin_steps = [
    'Login Admin — Masuk dengan akun admin',
    'Kelola Booking — Konfirmasi, batalkan, ubah status booking',
    'Kelola Layanan — Tambah/edit/hapus layanan & harga',
    'Kelola Barber — Tambah/edit/hapus barber, jadwal tidak tersedia',
    'Transaksi — Lihat riwayat pembayaran, filter by status/tanggal',
    'Laporan Pendapatan — Lihat rekap pendapatan per tanggal',
]

for i, step in enumerate(admin_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

# ===== SAVE =====
output_path = r'D:\FILE IRYA\SEMESTER 8\SKRIPSI\Semhas\arfan_barbershop\HAKI_Arfan_Barbershop.docx'
doc.save(output_path)
print(f'File saved to: {output_path}')